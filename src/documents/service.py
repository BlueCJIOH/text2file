from abc import ABC, abstractmethod
from io import BytesIO

import markdown
from docx import Document
from xhtml2pdf import pisa

from .constants import FileExtension


class DocumentGenerator(ABC):
    """
    Abstract base class for document generators.
    """

    @staticmethod
    @abstractmethod
    def generate(text: str) -> bytes:
        """
        Generates a document from the given markdown text.
        """
        pass

    @property
    @abstractmethod
    def media_type(self) -> str:
        """
        Returns the media type for the generated document.
        """
        pass

    @property
    @abstractmethod
    def filename_extension(self) -> str:
        """
        Returns the file extension for the generated document.
        """
        pass

class PDFGenerator(DocumentGenerator):
    """
    Concrete implementation for generating PDF documents.
    """

    @staticmethod
    def generate(text: str) -> bytes:
        """
        Generates a PDF document from markdown text.
        """

        # Convert markdown to HTML
        html = markdown.markdown(text)

        # Convert HTML to PDF
        pdf = BytesIO()
        pisa_status = pisa.CreatePDF(html, dest=pdf)
        if pisa_status.err:
            raise Exception("Error generating PDF document.")

        return pdf.getvalue()

    @property
    def media_type(self) -> str:
        return "application/pdf"

    @property
    def filename_extension(self) -> str:
        return "pdf"

class DOCXGenerator:
    """
    Concrete implementation for generating DOCX documents.
    """
    @staticmethod
    def generate(text: str) -> bytes:
        """
        Generates a DOCX document from markdown text.
        """

        # Convert markdown to plain text (no need for intermediate HTML)
        lines = text.split("\n")
        document = Document()

        for line in lines:
            # Detect Markdown-specific elements
            if line.startswith("# "):  # H1
                document.add_heading(line[2:], level=1)
            elif line.startswith("## "):  # H2
                document.add_heading(line[3:], level=2)
            elif line.startswith("### "):  # H3
                document.add_heading(line[4:], level=3)
            else:
                # Regular paragraph
                document.add_paragraph(line)

        # Save DOCX to bytes
        docx_file = BytesIO()
        document.save(docx_file)
        docx_file.seek(0)  # Reset the buffer position
        return docx_file.getvalue()

    @property
    def media_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    @property
    def filename_extension(self) -> str:
        return "docx"


class TXTGenerator(DocumentGenerator):
    """
    Concrete implementation for generating TEXT documents
    """

    @staticmethod
    def generate(text: str) -> bytes:
        """
        Generates a TEXT file without any markdown
        """
        return text.encode("utf-8")

    @property
    def media_type(self) -> str:
        return "text/plain"

    @property
    def filename_extension(self) -> str:
        return "txt"



class DocumentGeneratorFactory:
    """
    Factory class to get the appropriate document generator.
    """

    # Mapping extensions to their corresponding generator classes
    _generators = {
        FileExtension.PDF: PDFGenerator,
        FileExtension.DOCX: DOCXGenerator,
        FileExtension.TXT: TXTGenerator,
    }

    @staticmethod
    def get_generator(ext: FileExtension) -> DocumentGenerator:
        """
        Returns a document generator based on the file extension.
        """

        generator_class = DocumentGeneratorFactory._generators.get(ext)
        if not generator_class:
            raise ValueError(f"Unsupported file extension: {ext}")
        return generator_class()